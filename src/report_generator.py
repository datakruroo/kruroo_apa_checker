"""
report_generator.py
สร้าง Word document (.docx) จาก JSON results ของ apa_checker
มี 2 ส่วน: (1) สรุปผลการตรวจสอบ (2) ข้อเสนอแนะการปรับแก้
"""

import re
from datetime import datetime
from pathlib import Path
from collections import Counter
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .format_validation import sanitize_text_field
from .researcher_report_writer import polish_rows_with_llm, rows_from_results, user_level_label


# ===== สีที่ใช้ =====
COLOR_DARK = RGBColor(0x1A, 0x1A, 0x2E)     # navy dark
COLOR_PRIMARY = RGBColor(0x16, 0x21, 0x3E)   # dark blue
COLOR_ACCENT = RGBColor(0xE9, 0x4F, 0x37)    # red-orange (error)
COLOR_OK = RGBColor(0x06, 0x9A, 0x2E)        # green (ok)
COLOR_HEADER_BG = RGBColor(0x2C, 0x3E, 0x50) # table header bg
COLOR_ROW_ALT = RGBColor(0xF5, 0xF5, 0xF5)   # alternating row


SUMMARY_LABELS = [
    ("total_references", "จำนวนรายการอ้างอิงทั้งหมด"),
    ("verified_no_change_needed", "ตรวจยืนยันได้และยังไม่ต้องแก้"),
    ("verified_with_low_risk_formatting_fix", "มีการแก้เล็กน้อยที่ปลอดภัย"),
    ("bibliographic_conflicts", "ข้อมูลบรรณานุกรมขัดแย้ง ต้องตรวจเอง"),
    ("possible_matches_requiring_review", "พบรายการที่อาจเกี่ยวข้อง แต่ยังไม่ยืนยัน"),
    ("parser_warnings", "ข้อความที่อ่านจากไฟล์อาจเพี้ยน"),
    ("unsafe_generated_outputs", "ระบบไม่แสดงรายการที่สร้างใหม่ เพราะเสี่ยงผิด"),
    ("unverified_references", "ยังตรวจยืนยันไม่ได้"),
]

PROBLEM_LABELS = {
    "formatted output ถูก block เพราะไม่ผ่าน validation จึงไม่แสดงรายการสำหรับคัดลอก": "ยังไม่แสดงรายการสำหรับคัดลอก เพราะระบบอาจทำให้ข้อมูลเดิมเสีย",
    "formatted output ถูก block เพราะไม่ผ่าน validation จึงไม่แสดง canonical reference สำหรับคัดลอก": "ยังไม่แสดงรายการสำหรับคัดลอก เพราะระบบอาจทำให้ข้อมูลเดิมเสีย",
    "title_truncated": "ชื่อเรื่องที่ค้นได้ไม่ครบ หรือ subtitle อาจหาย",
    "missing_required_authors": "ข้อมูลผู้แต่งไม่ครบ จึงไม่เสนอให้แก้ผู้แต่ง",
    "missing_required_year": "ข้อมูลปีพิมพ์ไม่ครบ",
    "missing_required_title": "ข้อมูลชื่อเรื่องไม่ครบ",
    "missing_required_container_title": "ข้อมูลชื่อวารสารหรือแหล่งพิมพ์ไม่ครบ",
    "missing_required_book_title": "ข้อมูลชื่อหนังสือไม่ครบ",
    "missing_required_series_or_publisher": "ข้อมูลชุดรายงานหรือหน่วยงานเผยแพร่ไม่ครบ",
    "missing_required_proceedings_title": "ข้อมูลชื่อ proceedings ไม่ครบ",
    "html_entity_not_decoded": "มีอักขระที่ควรแปลงเป็นตัวอักษรปกติ เช่น &amp; เป็น &",
    "duplicate_terminal_punctuation": "มีจุดเกินหลังเครื่องหมายคำถามหรือเครื่องหมายตกใจ เช่น ?. หรือ !.",
    "midword_case_corruption": "พบตัวพิมพ์ใหญ่แทรกกลางคำผิดปกติ",
    "missing_year": "อ่านปีพิมพ์ไม่ได้จากรูปแบบรายการอ้างอิง",
    "doi_integrity": "DOI ในรายการที่สร้างไม่ตรงกับ DOI ต้นฉบับ",
    "unicode_dash_in_doi": "DOI มี dash แบบ Unicode ซึ่งไม่ควรเกิดขึ้น",
    "unicode_dash_in_url": "URL มี dash แบบ Unicode ซึ่งไม่ควรเกิดขึ้น",
    "url_integrity": "URL ในรายการที่สร้างไม่ตรงกับ URL ต้นฉบับ",
    "dropped_editors": "ข้อมูลบรรณาธิการหายไปจากรายการที่ระบบสร้าง",
    "dropped_publisher": "ข้อมูลสำนักพิมพ์หรือหน่วยงานเผยแพร่หายไปจากรายการที่ระบบสร้าง",
    "metadata-generated canonical reference ถูกปิดไว้; แสดงเฉพาะ safe fixes ที่ตรวจได้เท่านั้น": "ระบบไม่สร้างรายการอ้างอิงใหม่จากข้อมูลค้นหา และจะแสดงเฉพาะการแก้รูปแบบที่ปลอดภัย",
    "safe_fix:normalize_page_range": "แก้ช่วงหน้าให้ใช้ en dash (–) แทน hyphen (-) เฉพาะช่องหน้า",
    "safe_fix:normalize_doi_wrapper": "ปรับรูปแบบ DOI ให้ขึ้นต้นด้วย https://doi.org/ โดยไม่แก้ suffix ของ DOI",
    "safe_fix:sanitize_text": "แปลงอักขระที่อ่านมาเป็นรูปแบบปกติ เช่น &amp; เป็น &",
    "safe_fix:collapse_whitespace": "ลดช่องว่างซ้ำให้เหลือช่องว่างเดียว",
    "parser_warning": "ข้อความที่อ่านจากไฟล์อาจเพี้ยน",
    "listed_but_not_cited": "รายการอ้างอิงนี้อาจยังไม่ถูกอ้างในเนื้อหา",
    "ยืนยันตัวตน record ได้ด้วย exact DOI แต่ output ต้องผ่าน post-format validation แยกต่างหาก": "ตรวจพบ DOI ตรงกัน แต่ยังไม่ให้ระบบเขียนรายการใหม่แทนต้นฉบับ",
    "พบ candidate จาก title search แต่ยังไม่ใช่ exact DOI match จึงไม่แสดงเป็น verified": "ยังยืนยันจากแหล่งภายนอกไม่ได้ จึงตรวจได้เฉพาะรูปแบบ",
    "ยังไม่มีหลักฐาน metadata เพียงพอ ตรวจได้เฉพาะรูปแบบ APA ความเสี่ยงต่ำ": "ยังยืนยันจากแหล่งภายนอกไม่ได้ จึงตรวจได้เฉพาะรูปแบบ",
    "ไม่สามารถยืนยัน DOI ได้ จึงไม่ค้นหา/แทนที่ด้วย fuzzy search result อื่น": "ตรวจ DOI ไม่สำเร็จ และระบบจะไม่เดาแทนด้วยบทความอื่น",
    "The parsed text appears to be missing an inline token or formula.": "ข้อความที่อ่านจากไฟล์อาจขาดบางส่วน เช่น สูตร สัญลักษณ์ หรือคำกลางประโยค",
    "A word may have been split by a line-break hyphen.": "คำบางคำอาจถูกตัดด้วยเครื่องหมายยัติภังค์จากการอ่านไฟล์",
}


def _human_problem(text: str) -> str:
    return PROBLEM_LABELS.get(str(text), str(text))


def _human_status_lines(issue: dict) -> list[str]:
    lines = []
    severity = issue.get("severity")
    status = issue.get("status")
    identity = issue.get("identity_status")
    metadata = issue.get("metadata_status")
    formatting = issue.get("formatting_status")
    action = issue.get("action")

    if identity == "verified_exact_doi":
        lines.append("ยืนยัน DOI แล้ว")
    elif identity == "possible_match":
        lines.append("ยังยืนยันจากแหล่งภายนอกไม่ได้")
    elif identity == "conflicting":
        lines.append("ข้อมูลบรรณานุกรมขัดแย้ง ต้องตรวจเอง")
    elif identity == "unverified":
        lines.append("ยังตรวจยืนยันรายการนี้ไม่ได้")

    if metadata == "complete":
        lines.append("ข้อมูลจากแหล่งภายนอกมีข้อมูลขั้นต่ำครบ")
    elif metadata == "incomplete":
        lines.append("ข้อมูลจากแหล่งภายนอกยังไม่ครบ")
    elif metadata == "conflicting":
        lines.append("ข้อมูลจากแหล่งภายนอกขัดแย้งกับรายการต้นฉบับ")

    if formatting == "valid":
        lines.append("รูปแบบที่ตรวจได้ไม่พบความเสี่ยงสำคัญ")
    elif formatting == "unsafe_output":
        lines.append("รายการที่ระบบสร้างยังไม่ปลอดภัยพอให้คัดลอก")
    elif formatting == "parser_warning":
        lines.append("ควรตรวจไฟล์ต้นฉบับ เพราะข้อความที่อ่านมาอาจเพี้ยน")

    if action == "blocked":
        lines.append("ห้ามใช้รายการที่ระบบสร้างอัตโนมัติ")
    elif action == "human_review_required":
        lines.append("ต้องให้มนุษย์ตรวจรายการนี้")
    elif action == "auto_fix_safe":
        lines.append("แก้อัตโนมัติได้เฉพาะจุดเล็กน้อยที่ระบุ")
    elif action == "no_change_needed":
        lines.append("ยังไม่พบสิ่งที่ต้องแก้จากหลักฐานที่มี")

    if not lines and (severity == "PARSER_WARNING" or status == "parser_warning"):
        lines.append("ควรตรวจไฟล์ต้นฉบับ เพราะข้อความที่อ่านมาอาจเพี้ยน")

    return lines


def _human_evidence(item: str) -> str:
    text = str(item)
    match = re.search(r"Exact DOI match from ([^:]+):\s*(.+)", text)
    if match:
        return f"พบ DOI ตรงกันใน {_source_label(match.group(1))}: {match.group(2)}"
    if text == "Candidate from openalex.":
        return "พบรายการที่ชื่อคล้ายกันจาก OpenAlex แต่ยังไม่ยืนยันว่าเป็นรายการเดียวกัน"
    if text == "Candidate from crossref.":
        return "พบรายการที่ชื่อคล้ายกันจาก Crossref แต่ยังไม่ยืนยันว่าเป็นรายการเดียวกัน"
    if text == "Low-confidence candidate suppressed from report; debug only.":
        return "มีผลค้นหาคล้ายกันน้อยเกินไป จึงไม่แสดงเป็นข้อเสนอ"
    if text == "No DOI and no sufficiently attributable metadata candidate.":
        return "ไม่มี DOI และยังไม่พบข้อมูลจากแหล่งภายนอกที่ยืนยันได้"
    return _human_problem(text)


def _source_label(source: str) -> str:
    labels = {
        "crossref": "Crossref",
        "openalex": "OpenAlex",
        "manual": "ข้อมูลที่ป้อนเอง",
    }
    return labels.get(str(source).lower(), source)


def _format_confidence(value) -> str:
    try:
        return f"{float(value) * 100:.0f}%"
    except (TypeError, ValueError):
        return "-"


def _set_cell_bg(cell, hex_color: str):
    """ตั้งสีพื้นหลัง cell"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_heading(doc: Document, text: str, level: int = 1):
    """เพิ่ม heading ด้วย font ที่กำหนด"""
    para = doc.add_heading(text, level=level)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in para.runs:
        run.font.color.rgb = COLOR_PRIMARY
    return para


def _add_runs_from_markdown(para, text: str, color: RGBColor = None, size: Pt = Pt(10)):
    """
    แปลง *italic* Markdown เป็น Word italic จริงๆ
    ข้อความที่ไม่มี * → plain run
    ข้อความที่อยู่ใน *...* → italic run
    """
    text = sanitize_text_field(str(text))
    parts = re.split(r"\*([^*]+)\*", text)
    # parts สลับกัน: [plain, italic, plain, italic, ...]
    for i, part in enumerate(parts):
        if not part:
            continue
        run = para.add_run(part)
        run.font.size = size
        if color:
            run.font.color.rgb = color
        if i % 2 == 1:  # index คี่ = อยู่ใน *...*
            run.italic = True


def _add_paragraph(doc: Document, text: str, bold: bool = False, color: RGBColor = None):
    para = doc.add_paragraph()
    run = para.add_run(sanitize_text_field(text))
    run.font.size = Pt(11)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return para


def _build_summary_table(doc: Document, ref_check: dict, intext_check: dict):
    """สร้างตารางสรุปผลการตรวจสอบ"""
    doc.add_paragraph()  # spacing

    summary = ref_check.get("summary", {}) if ref_check else {}
    category_keys = SUMMARY_LABELS[1:]
    if any(key in summary for key, _label in category_keys):
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_row = table.rows[0]
        for i, h in enumerate(["หมวดผลการตรวจ", "จำนวน"]):
            cell = hdr_row.cells[i]
            _set_cell_bg(cell, "2C3E50")
            p = cell.paragraphs[0]
            run = p.add_run(h)
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(10)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rows = [(SUMMARY_LABELS[0][1], summary.get("total_references", 0))]
        rows.extend((label, summary.get(key, 0)) for key, label in category_keys)
        for label, value in rows:
            para = doc.add_paragraph()
            para.add_run(f"{label}: ").bold = True
            para.add_run(str(value)).font.size = Pt(10)
            row = table.add_row()
            row.cells[0].paragraphs[0].add_run(label).font.size = Pt(10)
            row.cells[1].paragraphs[0].add_run(str(value)).font.size = Pt(10)
            row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        return

    # header row + 2 data rows
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Column widths
    col_widths = [Cm(5), Cm(3.5), Cm(3.5), Cm(4)]
    for i, col in enumerate(table.columns):
        for cell in col.cells:
            cell.width = col_widths[i]

    # Header
    headers = ["หัวข้อการตรวจสอบ", "จำนวนทั้งหมด", "พบปัญหา", "ผ่าน"]
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        _set_cell_bg(cell, "2C3E50")
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # รายการอ้างอิง (References)
    if ref_check:
        total_ref = ref_check.get("summary", {}).get("total_references", 0)
        issues_ref = ref_check.get("summary", {}).get("issues_found", 0)
        ok_ref = total_ref - issues_ref
        row = table.add_row()
        data = ["รายการอ้างอิง (References)", str(total_ref), str(issues_ref), str(ok_ref)]
        for i, val in enumerate(data):
            cell = row.cells[i]
            _set_cell_bg(cell, "F5F5F5")
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(10)
            if i == 2 and issues_ref > 0:
                run.font.color.rgb = COLOR_ACCENT
            elif i == 3 and ok_ref == total_ref:
                run.font.color.rgb = COLOR_OK
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT

    # การอ้างอิงในเนื้อหา (In-Text)
    if intext_check:
        total_it = intext_check.get("summary", {}).get("total_checked", 0)
        issues_it = intext_check.get("summary", {}).get("issues_found", 0)
        ok_it = total_it - issues_it
        row = table.add_row()
        data = ["การอ้างอิงในเนื้อหา (In-Text)", str(total_it), str(issues_it), str(ok_it)]
        for i, val in enumerate(data):
            cell = row.cells[i]
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(10)
            if i == 2 and issues_it > 0:
                run.font.color.rgb = COLOR_ACCENT
            elif i == 3 and ok_it == total_it:
                run.font.color.rgb = COLOR_OK
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph()


def _build_researcher_table(
    doc: Document,
    results: dict,
    *,
    report_writer_client=None,
    report_writer_model: str = "",
):
    rows = rows_from_results(results)
    if report_writer_client and report_writer_model:
        rows = polish_rows_with_llm(report_writer_client, rows, report_writer_model)

    if not rows:
        _add_paragraph(
            doc,
            "ยังไม่พบรายการที่แก้รูปแบบได้อย่างปลอดภัยจากผลตรวจรอบนี้",
            color=COLOR_OK,
        )
        _add_paragraph(
            doc,
            "รายการที่ยังยืนยันข้อมูลไม่ได้หรือยังเสี่ยงผิดจะไม่ถูกเสนอเป็นข้อความให้คัดลอกในรายงานผู้วิจัย",
        )
        return

    intro = doc.add_paragraph()
    intro.add_run(
        "ส่วนนี้แสดงข้อเสนอแก้ไขที่ระบบตรวจได้จากรายการเดิม เช่น DOI ช่วงหน้า และรูปแบบตัวเอียง "
        "ระบบจะไม่แก้ผู้แต่ง ปี ชื่อเรื่อง หรือแหล่งพิมพ์โดยอัตโนมัติ หากพบข้อมูลจาก DOI หรือแหล่งภายนอก "
        "ไม่ตรงกับรายการเดิม ระบบจะแสดงคำเตือนเพื่อให้ผู้วิจัยตรวจสอบก่อนแก้ไข"
    ).font.size = Pt(10)

    note = doc.add_paragraph()
    note.add_run(
        "ความสำคัญของแต่ละรายการแสดงประเด็นที่ควรจัดการก่อน รายละเอียดภายในหนึ่งรายการอาจมีหลายประเด็นร่วมกัน"
    ).font.size = Pt(10)

    _build_researcher_level_summary(doc, rows)

    _build_researcher_issue_table(doc, rows)
    doc.add_paragraph()


def _build_researcher_level_summary(doc: Document, rows):
    counts = Counter(row.severity_level for row in rows)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, header in enumerate(["ระดับสูงสุดของรายการ", "จำนวนรายการ"]):
        cell = table.rows[0].cells[idx]
        _set_cell_bg(cell, "2C3E50")
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for level in ["ERROR", "WARNING", "STYLE_FIX", "AUTO_FIX"]:
        row = table.add_row()
        row.cells[0].paragraphs[0].add_run(user_level_label(level)).font.size = Pt(10)
        row.cells[1].paragraphs[0].add_run(str(counts.get(level, 0))).font.size = Pt(10)
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.add_run(f"รวมทั้งหมด {len(rows)} รายการที่ควรตรวจหรือแก้").font.size = Pt(10)
    p = doc.add_paragraph()
    p.add_run("ควรเริ่มจากรายการที่ต้องแก้ก่อนส่งบทความ ส่วนรายการที่ให้ตรวจสอบก่อนแก้ควรเทียบกับแหล่งเดิมก่อนแก้ข้อมูลสำคัญ").font.size = Pt(10)
    doc.add_paragraph()


def _build_researcher_issue_table(doc: Document, rows):
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["ความสำคัญ/รายการ", "เนื้อหาเดิม", "ประเด็นและข้อเสนอ", "ควรทำอย่างไร"]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        _set_cell_bg(cell, "2C3E50")
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row_data in rows:
        row = table.add_row()
        cells = row.cells
        level_text = f"{user_level_label(row_data.severity_level)}\n{row_data.item_label}"
        issue_text = row_data.proposed_text
        if row_data.reason:
            issue_text += f"\n\nเหตุผล:\n{row_data.reason}"
        values = [level_text, row_data.original_text, issue_text, row_data.next_action]
        for idx, value in enumerate(values):
            p = cells[idx].paragraphs[0]
            _add_runs_from_markdown(p, value, size=Pt(8))
            if idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_labeled_markdown_paragraph(doc: Document, label: str, text: str):
    if not text:
        return
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(10)
    _add_runs_from_markdown(p, text, size=Pt(10))


def _add_researcher_card(doc: Document, row_data):
    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(8)
    heading.paragraph_format.space_after = Pt(2)
    title = heading.add_run(f"{row_data.item_label} — ความสำคัญ: {user_level_label(row_data.severity_level)}")
    title.bold = True
    title.font.color.rgb = COLOR_PRIMARY
    title.font.size = Pt(12)

    _add_labeled_markdown_paragraph(doc, "เนื้อหาเดิม: ", row_data.original_text)
    _add_labeled_markdown_paragraph(doc, "ปัญหาที่พบ/ข้อเสนอให้ปรับแก้: ", row_data.proposed_text)
    _add_labeled_markdown_paragraph(doc, "เหตุผล: ", row_data.reason)
    _add_labeled_markdown_paragraph(doc, "ควรทำอย่างไร: ", row_data.next_action)

    divider = doc.add_paragraph()
    divider.add_run("─" * 48).font.size = Pt(8)


def _build_issue_types_list(doc: Document, ref_check: dict, intext_check: dict):
    """แสดงรายการประเภทปัญหาที่พบ"""
    issue_types = []
    if ref_check:
        issue_types += ref_check.get("summary", {}).get("issue_types", [])
    if intext_check:
        issue_types += intext_check.get("summary", {}).get("issue_types", [])

    if issue_types:
        _add_paragraph(doc, "สิ่งที่ควรระวัง:", bold=True)
        for it in sorted({_human_problem(it) for it in issue_types}):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(it).font.size = Pt(11)
    doc.add_paragraph()


def _build_reference_issues(doc: Document, ref_check: dict):
    """สร้างส่วนรายละเอียดปัญหา references"""
    if not ref_check:
        return

    issues = ref_check.get("issues", [])
    if not issues:
        p = _add_paragraph(doc, "✓ ไม่พบข้อผิดพลาดในรายการอ้างอิง", color=COLOR_OK)
        return

    # กรองเฉพาะรายการที่มีปัญหาจริง (issues list ไม่ว่าง)
    real_issues = [iss for iss in issues if iss.get("issues")]
    if not real_issues:
        _add_paragraph(doc, "✓ ไม่พบข้อผิดพลาดในรายการอ้างอิง", color=COLOR_OK)
        return

    for issue in real_issues:
        num = issue.get("reference_number", "?")
        _add_paragraph(doc, f"รายการที่ {num}", bold=True, color=COLOR_PRIMARY)

        # ต้นฉบับ (แปลง *italic* → Word italic จริง)
        p = doc.add_paragraph()
        p.add_run("ต้นฉบับ:  ").bold = True
        _add_runs_from_markdown(p, issue.get("original", ""), color=COLOR_ACCENT)

        # ปัญหาที่พบ
        for prob in [_human_problem(prob) for prob in issue.get("issues", [])]:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(prob)
            run.font.size = Pt(10)
            run.font.color.rgb = COLOR_ACCENT

        if issue.get("severity") or issue.get("status"):
            p = doc.add_paragraph()
            p.add_run("ผลตรวจ: ").bold = True
            p.add_run("; ".join(_human_status_lines(issue))).font.size = Pt(10)

        if issue.get("evidence"):
            p = doc.add_paragraph()
            p.add_run("หลักฐาน: ").bold = True
            p.add_run("; ".join(_human_evidence(item) for item in issue.get("evidence", []) if item)).font.size = Pt(10)

        metadata_bits = []
        if issue.get("metadata_source"):
            metadata_bits.append(f"แหล่งข้อมูล: {_source_label(issue.get('metadata_source'))}")
        if issue.get("source_record_id"):
            metadata_bits.append(f"รหัสรายการ: {issue.get('source_record_id')}")
        if issue.get("doi"):
            metadata_bits.append(f"DOI: {issue.get('doi')}")
        if issue.get("confidence") is not None:
            metadata_bits.append(f"ความมั่นใจของการจับคู่: {_format_confidence(issue.get('confidence'))}")
        if metadata_bits:
            p = doc.add_paragraph()
            p.add_run("ข้อมูลประกอบการตรวจ: ").bold = True
            p.add_run(" | ".join(metadata_bits)).font.size = Pt(10)

        if issue.get("auto_fixable") is not None:
            p = doc.add_paragraph()
            p.add_run("การแก้อัตโนมัติ: ").bold = True
            if issue.get("auto_fixable"):
                auto_text = "อนุญาตเฉพาะความเสี่ยงต่ำ"
            elif issue.get("action") == "no_change_needed":
                auto_text = "ไม่ต้องแก้"
            elif issue.get("action") == "blocked":
                auto_text = "ไม่อนุญาต เพราะรายการที่ระบบสร้างยังเสี่ยงผิด"
            elif issue.get("action") == "human_review_required":
                auto_text = "ต้องให้มนุษย์ตรวจ"
            else:
                auto_text = "ต้องให้มนุษย์ตรวจ"
            p.add_run(auto_text).font.size = Pt(10)

        # ข้อเสนอ correction แยก verified exact metadata ออกจาก proposal ทั่วไป
        if issue.get("corrected"):
            p = doc.add_paragraph()
            p.add_run(f"{issue.get('suggestion_label', 'ข้อเสนอเบื้องต้น')}: ").bold = True
            _add_runs_from_markdown(p, issue.get("corrected", ""), color=COLOR_OK)
        else:
            p = doc.add_paragraph()
            p.add_run("ข้อเสนอเบื้องต้น: ").bold = True
            if issue.get("action") == "no_change_needed":
                proposal_text = "ไม่มีข้อเสนอให้คัดลอก เพราะยังไม่พบจุดที่ต้องแก้แบบปลอดภัย"
            else:
                proposal_text = "ยังไม่แสดงรายการสำหรับคัดลอก เพราะยังยืนยันข้อมูลไม่ได้หรือรายการที่สร้างยังเสี่ยงผิด"
            p.add_run(proposal_text).font.size = Pt(10)

        # คำอธิบาย
        if issue.get("explanation"):
            p = doc.add_paragraph()
            p.add_run("คำอธิบาย: ").bold = True
            p.add_run("ระบบตรวจตัวตนของแหล่งข้อมูล แยกจากการตรวจว่ารายการที่จัดรูปแบบแล้วปลอดภัยพอให้ใช้หรือไม่").font.size = Pt(10)

        doc.add_paragraph("─" * 60)


def _build_intext_issues(doc: Document, intext_check: dict):
    """สร้างส่วนรายละเอียดปัญหา in-text citations"""
    if not intext_check:
        return

    issues = intext_check.get("issues", [])
    if not issues:
        _add_paragraph(doc, "✓ ไม่พบข้อผิดพลาดในการอ้างอิงในเนื้อหา", color=COLOR_OK)
        return

    real_issues = [iss for iss in issues if iss.get("issues")]
    if not real_issues:
        _add_paragraph(doc, "✓ ไม่พบข้อผิดพลาดในการอ้างอิงในเนื้อหา", color=COLOR_OK)
        return

    for issue in real_issues:
        num = issue.get("excerpt_number", "?")
        _add_paragraph(doc, f"ประโยคที่ {num}", bold=True, color=COLOR_PRIMARY)

        p = doc.add_paragraph()
        p.add_run("ต้นฉบับ:  ").bold = True
        _add_runs_from_markdown(p, f'"{issue.get("excerpt", "")}"')

        for prob in issue.get("issues", []):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(prob)
            run.font.size = Pt(10)
            run.font.color.rgb = COLOR_ACCENT

        if issue.get("corrected"):
            p = doc.add_paragraph()
            p.add_run(f"{issue.get('suggestion_label', 'ข้อเสนอเบื้องต้น')}: ").bold = True
            _add_runs_from_markdown(p, issue.get("corrected", ""), color=COLOR_OK)

        if issue.get("explanation"):
            p = doc.add_paragraph()
            p.add_run("คำอธิบาย: ").bold = True
            p.add_run(issue.get("explanation", "")).font.size = Pt(10)

        doc.add_paragraph("─" * 60)


def generate_report(
    results: dict,
    article_filename: str,
    output_path: str,
    checklist_name: str = "APA 7th Edition (จุฬาฯ 2568)",
    report_writer_client=None,
    report_writer_model: str = "",
    include_technical_section: bool = False,
) -> str:
    """
    สร้าง Word report จาก results ของ apa_checker.run_full_check()

    Args:
        results: dict จาก run_full_check()
        article_filename: ชื่อไฟล์บทความที่ตรวจ
        output_path: path สำหรับบันทึก .docx
        checklist_name: ชื่อเกณฑ์ที่ใช้ตรวจ

    Returns:
        path ของไฟล์ที่สร้าง
    """
    doc = Document()

    # ตั้งค่า margin
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    # ===== หน้าปก =====
    title_para = doc.add_heading("รายงานผลการตรวจสอบการอ้างอิง", 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(f"เกณฑ์: {checklist_name}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"ไฟล์บทความ: {article_filename}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        f"วันที่ตรวจสอบ: {datetime.now().strftime('%d %B %Y, %H:%M น.')}"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ===== ข้อเสนอแนะสำหรับผู้วิจัย =====
    _add_heading(doc, "ข้อเสนอแนะสำหรับผู้วิจัย", level=1)

    if not results.get("ref_found"):
        _add_paragraph(
            doc,
            "ไม่พบส่วนรายการอ้างอิงในเอกสาร กรุณาตรวจสอบว่าไฟล์มีหัวข้อรายการอ้างอิงที่อ่านได้",
            color=COLOR_ACCENT,
        )

    _build_researcher_table(
        doc,
        results,
        report_writer_client=report_writer_client,
        report_writer_model=report_writer_model,
    )

    if not include_technical_section:
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        return output_path

    doc.add_page_break()

    # ===== รายละเอียดเชิงเทคนิค =====
    _add_heading(doc, "รายละเอียดเชิงเทคนิค", level=1)

    ref_check = results.get("ref_check")
    intext_check = results.get("intext_check")

    if not results.get("ref_found"):
        _add_paragraph(
            doc,
            "⚠ ไม่พบส่วน 'รายการอ้างอิง' ในเอกสาร กรุณาตรวจสอบว่า PDF มีส่วนอ้างอิงที่อ่านได้",
            color=COLOR_ACCENT,
        )
    else:
        _build_summary_table(doc, ref_check, intext_check)
        _build_issue_types_list(doc, ref_check, intext_check)

    doc.add_page_break()

    _add_heading(doc, "2.1 รายการอ้างอิง (References)", level=2)
    _build_reference_issues(doc, ref_check)

    doc.add_paragraph()
    _add_heading(doc, "2.2 การอ้างอิงในเนื้อหา (In-Text Citation)", level=2)
    _build_intext_issues(doc, intext_check)

    # บันทึกไฟล์
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path
