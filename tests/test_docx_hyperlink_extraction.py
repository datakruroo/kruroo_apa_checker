import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from src.docx_extractor import _parse_docx_to_markdown


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


class DocxHyperlinkExtractionTests(unittest.TestCase):
    def test_docx_extractor_preserves_hyperlink_text_in_references(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "fixture.docx"
            doc = Document()
            doc.add_paragraph("References")
            paragraph = doc.add_paragraph(
                "Treffinger, D. J. (2008). A title. Journal, 18(4), 390-401. "
            )
            _add_hyperlink(
                paragraph,
                "https://doi.org/https://doi.org/10.1016/j.lindif.2007.11.007",
                "https://doi.org/https://doi.org/10.1016/j.lindif.2007.11.007",
            )
            doc.save(path)

            markdown = _parse_docx_to_markdown(str(path))

        self.assertIn("https://doi.org/https://doi.org/10.1016/j.lindif.2007.11.007", markdown)


if __name__ == "__main__":
    unittest.main()
