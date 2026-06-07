import tempfile
import unittest
from pathlib import Path

from docx import Document

from src.report_generator import generate_report
from src.researcher_report_writer import (
    ResearcherReportRow,
    polish_rows_with_llm_output,
    rows_from_results,
)


class ResearcherReportWriterTests(unittest.TestCase):
    def test_rows_from_results_use_plain_researcher_language(self):
        results = {
            "ref_check": {
                "issues": [
                    {
                        "reference_number": 14,
                        "original": "Siemens, G. (2013). Learning analytics: The emergence of a discipline.",
                        "corrected": "",
                        "identity_status": "verified_exact_doi",
                        "formatting_status": "unsafe_output",
                        "action": "blocked",
                        "issues": ["title_truncated"],
                    },
                    {
                        "reference_number": 2,
                        "original": "Author, A. (2020). A title. Journal, 1, 10-20.",
                        "corrected": "Author, A. (2020). A title. Journal, 1, 10–20.",
                        "action": "auto_fix_safe",
                        "issues": ["safe_fix:normalize_page_range"],
                    },
                ]
            }
        }

        rows = rows_from_results(results)

        self.assertEqual(len(rows), 1)
        self.assertEqual(rows[0].item_label, "รายการที่ 2")
        self.assertIn("10–20", rows[0].proposed_text)
        self.assertIn("ปรับรูปแบบช่วงหน้า", rows[0].reason)

    def test_unverified_reference_only_gets_format_feedback(self):
        results = {
            "ref_check": {
                "issues": [
                    {
                        "reference_number": 3,
                        "original": "Author, A. (2020). A title. Journal, 1, 10-20.",
                        "corrected": "Author, A. (2020). A title. Journal, 1, 10–20.",
                        "identity_status": "unverified",
                        "action": "auto_fix_safe",
                        "issues": ["safe_fix:normalize_page_range"],
                    },
                    {
                        "reference_number": 4,
                        "original": "Author, A. (2020). A title.",
                        "corrected": "",
                        "identity_status": "unverified",
                        "action": "human_review_required",
                        "issues": ["ยังไม่มีหลักฐาน metadata เพียงพอ ตรวจได้เฉพาะรูปแบบ APA ความเสี่ยงต่ำ"],
                    },
                ]
            }
        }

        rows = rows_from_results(results)

        self.assertEqual(len(rows), 1)
        self.assertIn("10–20", rows[0].proposed_text)
        self.assertIn("ปรับรูปแบบช่วงหน้า", rows[0].reason)

    def test_researcher_table_skips_unverified_metadata_only_rows(self):
        results = {
            "ref_check": {
                "issues": [
                    {
                        "reference_number": 4,
                        "original": "Author, A. (2020). A title.",
                        "corrected": "",
                        "identity_status": "unverified",
                        "action": "human_review_required",
                        "issues": ["ยังไม่มีหลักฐาน metadata เพียงพอ ตรวจได้เฉพาะรูปแบบ APA ความเสี่ยงต่ำ"],
                    },
                    {
                        "reference_number": 5,
                        "original": "Siemens, G. (2013). Learning analytics: The emergence of a discipline.",
                        "corrected": "",
                        "identity_status": "verified_exact_doi",
                        "formatting_status": "unsafe_output",
                        "action": "blocked",
                        "issues": ["title_truncated"],
                    },
                ]
            }
        }

        rows = rows_from_results(results)

        self.assertEqual(rows, [])

    def test_researcher_table_shows_safe_fix_details_instead_of_full_reference_when_available(self):
        results = {
            "ref_check": {
                "issues": [
                    {
                        "reference_number": 8,
                        "original": "Brown, T. (2008). Design thinking. Harvard Business Review, 86(6), 84-92.",
                        "corrected": "Brown, T. (2008). Design thinking. Harvard Business Review, 86(6), 84–92.",
                        "action": "auto_fix_safe",
                        "issues": ["safe_fix:normalize_page_range"],
                        "safe_fix_details": ["ช่วงหน้า: 84-92 → 84–92"],
                    }
                ]
            }
        }

        rows = rows_from_results(results)

        self.assertEqual(len(rows), 1)
        self.assertIn("แก้เฉพาะจุดนี้", rows[0].proposed_text)
        self.assertIn("ช่วงหน้า: 84-92 → 84–92", rows[0].proposed_text)
        self.assertNotIn("Harvard Business Review", rows[0].proposed_text)

    def test_researcher_table_skips_intext_without_proposed_fix(self):
        results = {
            "intext_check": {
                "issues": [
                    {
                        "excerpt_number": 2,
                        "excerpt": "(Author, 2020; Other, 2021)",
                        "corrected": "",
                        "issues": ["ต้องตรวจรายการนี้เพิ่มเติม"],
                    }
                ]
            }
        }

        rows = rows_from_results(results)

        self.assertEqual(rows, [])

    def test_researcher_table_hides_possible_unmatched_citation_until_precise(self):
        results = {
            "intext_check": {
                "issues": [
                    {
                        "excerpt_number": 1,
                        "excerpt": "(Idin, 2020)",
                        "corrected": "",
                        "issues": ["possible_unmatched_citation"],
                        "format_finding_details": ["พบ citation ที่ยังจับคู่กับ References ไม่ได้"],
                    }
                ]
            }
        }

        rows = rows_from_results(results)

        self.assertEqual(rows, [])

    def test_researcher_table_includes_citation_reference_year_mismatch(self):
        results = {
            "intext_check": {
                "issues": [
                    {
                        "excerpt_number": 3,
                        "excerpt": "(กวินนาฏ พลอยกระจ่าง และคณะ, 2564)",
                        "corrected": "",
                        "issues": ["possible_unmatched_citation", "possible_year_mismatch"],
                        "format_finding_details": [
                            "พบ citation ที่ยังจับคู่กับ References ไม่ได้",
                            "พบ citation ปี 2564 แต่ References ของผู้แต่งนี้มีปี 2563",
                        ],
                    }
                ]
            }
        }

        rows = rows_from_results(results)

        self.assertEqual(len(rows), 1)
        self.assertEqual(rows[0].item_label, "ประโยคที่ 3")
        self.assertEqual(rows[0].severity_level, "ERROR")
        self.assertIn("2564", rows[0].proposed_text)
        self.assertNotIn("จับคู่กับ References ไม่ได้", rows[0].proposed_text)
        self.assertNotIn("จับคู่กับ References ไม่ได้", rows[0].reason)
        self.assertIn("2563", rows[0].reason)
        self.assertIn("จับคู่กับรายการอ้างอิง", rows[0].next_action)

    def test_researcher_table_includes_uncited_reference_warning(self):
        results = {
            "intext_check": {
                "issues": [
                    {
                        "reference_number": 8,
                        "excerpt": "Unused, B. (2021). Another title.",
                        "corrected": "",
                        "issues": ["listed_but_not_cited"],
                        "format_finding_details": ["ยังไม่พบ citation ในเนื้อหาที่ตรงกับ Unused, B. (2021)"],
                    }
                ]
            }
        }

        rows = rows_from_results(results)

        self.assertEqual(len(rows), 1)
        self.assertEqual(rows[0].severity_level, "WARNING")
        self.assertEqual(rows[0].item_label, "รายการอ้างอิงที่ 8")
        self.assertIn("ยังไม่พบ citation", rows[0].reason)
        self.assertIn("พิจารณาลบเอง", rows[0].next_action)

    def test_uncited_warning_merges_into_existing_reference_card(self):
        results = {
            "ref_check": {
                "issues": [
                    {
                        "reference_number": 3,
                        "original": "Author Name (2020). A title.",
                        "issues": ["format:apa_author_period_before_year"],
                        "format_findings": [
                            {
                                "rule_id": "apa_author_period_before_year",
                                "field": "author",
                                "message": "หลังชื่อผู้แต่งก่อนปีพิมพ์ควรมีจุด",
                                "detail": "ชื่อผู้แต่ง/ปี: Author Name (2020) → Author Name. (2020)",
                                "safe_to_apply": True,
                            }
                        ],
                    }
                ]
            },
            "intext_check": {
                "issues": [
                    {
                        "reference_number": 3,
                        "excerpt": "Author Name (2020). A title.",
                        "corrected": "",
                        "issues": ["listed_but_not_cited"],
                        "format_finding_details": ["ยังไม่พบ citation ในเนื้อหาที่ตรงกับ Author Name (2020)"],
                    }
                ]
            },
        }

        rows = rows_from_results(results)

        self.assertEqual(len(rows), 1)
        self.assertEqual(rows[0].item_label, "รายการที่ 3")
        self.assertEqual(rows[0].severity_level, "WARNING")
        self.assertIn("แก้ได้เลย", rows[0].proposed_text)
        self.assertIn("ควรตรวจสอบก่อนแก้: ยังไม่พบ citation", rows[0].proposed_text)
        self.assertNotIn("[AUTO_FIX]", rows[0].proposed_text)
        self.assertNotIn("[WARNING]", rows[0].proposed_text)

    def test_llm_polish_cannot_change_original_or_proposed_text(self):
        rows = [
            ResearcherReportRow(
                severity_level="AUTO_FIX",
                item_label="รายการที่ 1",
                original_text="Original reference",
                proposed_text="Proposed safe fix",
                reason="เดิม",
                next_action="เดิม",
                technical_note="note",
            )
        ]
        llm_payload = {
            "rows": [
                {
                    "item_label": "รายการที่ 1",
                    "original_text": "CHANGED",
                    "proposed_text": "CHANGED",
                    "reason": "เหตุผลที่เกลาแล้ว",
                    "next_action": "คำแนะนำที่เกลาแล้ว",
                }
            ]
        }

        polished = polish_rows_with_llm_output(rows, llm_payload)

        self.assertEqual(polished[0].original_text, "Original reference")
        self.assertEqual(polished[0].proposed_text, "Proposed safe fix")
        self.assertEqual(polished[0].severity_level, "AUTO_FIX")
        self.assertEqual(polished[0].reason, "เหตุผลที่เกลาแล้ว")
        self.assertEqual(polished[0].next_action, "คำแนะนำที่เกลาแล้ว")

    def test_docx_report_defaults_to_researcher_section_only(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            output = Path(tmpdir) / "report.docx"
            generate_report(
                results={
                    "ref_found": True,
                    "ref_check": {
                        "summary": {
                            "total_references": 1,
                            "verified_no_change_needed": 0,
                            "verified_with_low_risk_formatting_fix": 0,
                            "bibliographic_conflicts": 0,
                            "possible_matches_requiring_review": 0,
                            "parser_warnings": 0,
                            "unsafe_generated_outputs": 1,
                            "unverified_references": 0,
                            "issues_found": 1,
                            "issue_types": ["title_truncated"],
                        },
                        "issues": [
                            {
                                "reference_number": 14,
                                "original": "Siemens, G. (2013). Learning analytics: The emergence of a discipline.",
                                "issues": ["title_truncated"],
                                "corrected": "",
                                "identity_status": "verified_exact_doi",
                                "metadata_status": "complete",
                                "formatting_status": "unsafe_output",
                                "action": "blocked",
                                "evidence": [],
                                "confidence": 1.0,
                            }
                        ],
                    },
                    "intext_check": {"summary": {"total_checked": 0, "issues_found": 0}, "issues": []},
                },
                article_filename="fixture.docx",
                output_path=str(output),
            )
            doc = Document(output)
            text = "\n".join(p.text for p in doc.paragraphs)

        self.assertIn("ข้อเสนอแนะสำหรับผู้วิจัย", text)
        self.assertNotIn("รายละเอียดเชิงเทคนิค", text)
        self.assertIn("ยังไม่พบรายการที่แก้รูปแบบได้อย่างปลอดภัย", text)

    def test_docx_report_uses_summary_and_cards_for_researcher_rows(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            output = Path(tmpdir) / "report.docx"
            generate_report(
                results={
                    "ref_found": True,
                    "ref_check": {
                        "summary": {"total_references": 1, "issues_found": 1},
                        "issues": [
                            {
                                "reference_number": 12,
                                "original": "Brown, T. (2008). Design thinking. Journal, 86(6), 84-92.",
                                "issues": ["format:apa_page_range_en_dash"],
                                "format_findings": [
                                    {
                                        "rule_id": "apa_page_range_en_dash",
                                        "field": "pages",
                                        "message": "ช่วงหน้าควรใช้ en dash",
                                        "detail": "ช่วงหน้า: 84-92 → 84–92",
                                        "safe_to_apply": True,
                                    }
                                ],
                            }
                        ],
                    },
                    "intext_check": {"summary": {"total_checked": 0, "issues_found": 0}, "issues": []},
                },
                article_filename="fixture.docx",
                output_path=str(output),
            )
            doc = Document(output)
            text = "\n".join(p.text for p in doc.paragraphs)
            table_count = len(doc.tables)
            table_text = "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)

        self.assertIn("ความสำคัญของแต่ละรายการแสดงประเด็นที่ควรจัดการก่อน", text)
        self.assertIn("ความสำคัญ/รายการ", table_text)
        self.assertIn("แก้ได้เลย\nรายการที่ 12", table_text)
        self.assertNotIn("AUTO_FIX\nรายการที่ 12", table_text)
        self.assertIn("ประเด็นและข้อเสนอ", table_text)
        self.assertIn("ระดับสูงสุดของรายการ", table_text)
        self.assertIn("จำนวนรายการ", table_text)
        self.assertIn("รวมทั้งหมด 1 รายการที่ควรตรวจหรือแก้", text)
        self.assertIn("ควรเริ่มจากรายการที่ต้องแก้ก่อนส่งบทความ", text)
        self.assertGreaterEqual(table_count, 2)

    def test_docx_report_sanitizes_text_before_render(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            output = Path(tmpdir) / "report.docx"
            generate_report(
                results={
                    "ref_found": True,
                    "ref_check": {
                        "summary": {"total_references": 1, "issues_found": 1},
                        "issues": [
                            {
                                "reference_number": 1,
                                "original": "Author\u200b, A. (2020). Soft\u00adhyphen title. Journal, 1, 1-2.",
                                "issues": ["format:apa_page_range_en_dash"],
                                "format_findings": [
                                    {
                                        "rule_id": "apa_page_range_en_dash",
                                        "field": "pages",
                                        "message": "ช่วงหน้าควรใช้ en dash (–) แทน hyphen (-)",
                                        "detail": "ช่วงหน้า: 1-2 → 1–2",
                                        "safe_to_apply": True,
                                    }
                                ],
                            }
                        ],
                    },
                    "intext_check": {"summary": {"total_checked": 0, "issues_found": 0}, "issues": []},
                },
                article_filename="fixture.docx",
                output_path=str(output),
            )
            doc = Document(output)
            text = "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)

        self.assertNotIn("\u200b", text)
        self.assertNotIn("\u00ad", text)
        self.assertIn("Softhyphen", text)

    def test_docx_report_can_include_technical_section_when_requested(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            output = Path(tmpdir) / "report.docx"
            generate_report(
                results={
                    "ref_found": True,
                    "ref_check": {
                        "summary": {"total_references": 1, "issues_found": 0},
                        "issues": [],
                    },
                    "intext_check": {"summary": {"total_checked": 0, "issues_found": 0}, "issues": []},
                },
                article_filename="fixture.docx",
                output_path=str(output),
                include_technical_section=True,
            )
            doc = Document(output)
            text = "\n".join(p.text for p in doc.paragraphs)

        self.assertIn("รายละเอียดเชิงเทคนิค", text)


if __name__ == "__main__":
    unittest.main()
